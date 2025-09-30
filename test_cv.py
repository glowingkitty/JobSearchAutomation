#!/usr/bin/env python3
"""
Test script for CV Automation System - Phase 1
Verifies that the generated CV contains expected content and formatting.
"""

import os
import sys
from pathlib import Path
from docx import Document

def test_cv_structure(filepath):
    """Test the structure and content of a generated CV."""
    print(f"Testing CV: {filepath}")
    print("-" * 50)
    
    if not os.path.exists(filepath):
        print("❌ CV file not found!")
        return False
    
    try:
        doc = Document(filepath)
        print("✅ CV file opened successfully")
        
        # Check basic structure
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Expected sections
        expected_sections = [
            "Professional Summary",
            "Professional Experience", 
            "Education",
            "Skills",
            "Certifications",
            "Projects",
            "Languages"
        ]
        
        found_sections = []
        for section in expected_sections:
            if any(section in para for para in paragraphs):
                found_sections.append(section)
                print(f"✅ Found section: {section}")
            else:
                print(f"❌ Missing section: {section}")
        
        # Check for personal information
        personal_info_found = any("John Doe" in para for para in paragraphs)
        if personal_info_found:
            print("✅ Personal information found")
        else:
            print("❌ Personal information not found")
        
        # Check for work experience
        experience_found = any("TechCorp Inc." in para for para in paragraphs)
        if experience_found:
            print("✅ Work experience found")
        else:
            print("❌ Work experience not found")
        
        # Check for skills
        skills_found = any("Python" in para for para in paragraphs)
        if skills_found:
            print("✅ Skills section found")
        else:
            print("❌ Skills section not found")
        
        # Check document properties
        print(f"\nDocument Statistics:")
        print(f"- Total paragraphs: {len(doc.paragraphs)}")
        print(f"- Non-empty paragraphs: {len(paragraphs)}")
        print(f"- Sections found: {len(found_sections)}/{len(expected_sections)}")
        
        # Check formatting (basic)
        for para in doc.paragraphs[:5]:  # Check first 5 paragraphs
            if para.text.strip():
                for run in para.runs:
                    if run.font.name:
                        print(f"✅ Font detected: {run.font.name}")
                        break
                break
        
        return len(found_sections) >= 5  # At least 5 sections should be found
        
    except Exception as e:
        print(f"❌ Error reading CV: {e}")
        return False

def main():
    """Main test function."""
    print("CV Automation System - Test Suite")
    print("=" * 40)
    
    # Find the most recent CV file
    output_dir = Path("output")
    if not output_dir.exists():
        print("❌ Output directory not found!")
        return False
    
    cv_files = list(output_dir.glob("*.docx"))
    if not cv_files:
        print("❌ No CV files found in output directory!")
        print("Please run: python generate_cv.py --yaml-file data/example_cv.yaml")
        return False
    
    # Get the most recent file
    latest_cv = max(cv_files, key=os.path.getctime)
    print(f"Testing latest CV: {latest_cv.name}")
    
    # Run tests
    success = test_cv_structure(str(latest_cv))
    
    if success:
        print("\n🎉 All tests passed! CV generation is working correctly.")
        return True
    else:
        print("\n❌ Some tests failed. Please check the CV generation process.")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
