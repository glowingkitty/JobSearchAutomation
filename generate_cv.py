#!/usr/bin/env python3
"""
CV Generation Script - Phase 1
Generates ATS-friendly CV documents from YAML data using python-docx.

This script reads personal and professional information from a YAML file
and generates a professional CV in DOCX format optimized for ATS systems.
"""

import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Configure logging first
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('cv_generation.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# PDF generation imports
try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
    logger.info("WeasyPrint available for PDF generation")
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logger.warning("WeasyPrint not available. PDF generation will be disabled.")

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
    logger.info("pdfkit available for PDF generation")
except ImportError:
    PDFKIT_AVAILABLE = False
    logger.warning("pdfkit not available. PDF generation will be disabled.")

# Try to import python-dotenv for environment variable loading
try:
    from dotenv import load_dotenv
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False
    logger.warning("python-dotenv not available. Environment variables must be set manually.")


class CVGenerator:
    """
    Generates professional CV documents from YAML data.
    Optimized for ATS (Applicant Tracking System) compatibility.
    """
    
    def __init__(self, yaml_file: str = "data/master_cv.yaml"):
        """
        Initialize CV generator with YAML data file.
        
        Args:
            yaml_file: Path to YAML file containing CV data
        """
        self.yaml_file = yaml_file
        self.data = None
        self.doc = None
        self.config = {}
        self.personal_info = {}
        
        logger.info(f"Initializing CV generator with file: {yaml_file}")
        
        # Load environment variables for personal information
        self.load_personal_info_from_env()
    
    def load_personal_info_from_env(self):
        """
        Load personal information from environment variables.
        First tries to load from personal_info.env file, then falls back to system env vars.
        """
        # Try to load from personal_info.env file first
        env_file = "personal_info.env"
        if os.path.exists(env_file):
            if DOTENV_AVAILABLE:
                load_dotenv(env_file)
                logger.info(f"Loaded environment variables from {env_file}")
            else:
                logger.warning(f"Found {env_file} but python-dotenv not available. Please install it or set environment variables manually.")
        else:
            logger.info(f"No {env_file} file found, using system environment variables")
        
        # Load personal information from environment variables
        self.personal_info = {
            'name': os.getenv('CV_NAME', ''),
            'email': os.getenv('CV_EMAIL', ''),
            'phone': os.getenv('CV_PHONE', ''),
            'location': os.getenv('CV_LOCATION', ''),
            'linkedin': os.getenv('CV_LINKEDIN', ''),
            'website': os.getenv('CV_WEBSITE', ''),
            'github': os.getenv('CV_GITHUB', '')
        }
        
        # Validate required fields
        required_fields = ['name', 'email']
        missing_fields = [field for field in required_fields if not self.personal_info[field]]
        
        if missing_fields:
            logger.error(f"Missing required personal information: {', '.join(missing_fields)}")
            logger.error("Please set the following environment variables:")
            for field in missing_fields:
                logger.error(f"  CV_{field.upper()}")
            logger.error("Or create a personal_info.env file with these values.")
        else:
            logger.info("Personal information loaded successfully from environment variables")
    
    def validate_yaml_structure(self) -> bool:
        """
        Validate the structure of the loaded YAML data.
        
        Returns:
            bool: True if structure is valid, False otherwise
        """
        errors = []
        
        # Check required sections (personal_info is now loaded from environment variables)
        required_sections = []
        for section in required_sections:
            if section not in self.data:
                errors.append(f"Missing required section: '{section}'")
        
        # Personal information is now loaded from environment variables
        # Validate that we have the required personal info from environment
        required_personal_fields = ['name', 'email']
        missing_personal_fields = [field for field in required_personal_fields if not self.personal_info.get(field)]
        if missing_personal_fields:
            errors.append(f"Missing required personal information from environment: {', '.join(missing_personal_fields)}")
        
        # Validate list sections (experience, education, projects, languages, certifications)
        list_sections = ['experience', 'education', 'projects', 'languages', 'certifications']
        for section in list_sections:
            if section in self.data:
                if self.data[section] is None:
                    # Empty section with just comments - this is OK
                    continue
                elif not isinstance(self.data[section], list):
                    errors.append(f"'{section}' must be a list, got {type(self.data[section]).__name__}")
                else:
                    # Check each entry is a dictionary
                    for i, item in enumerate(self.data[section]):
                        if not isinstance(item, dict):
                            errors.append(f"'{section}' entry {i+1} must be a dictionary, got {type(item).__name__}: {item}")
        
        # Validate skills structure
        if 'skills' in self.data:
            skills = self.data['skills']
            if not isinstance(skills, dict):
                errors.append("'skills' must be a dictionary with skill categories")
            else:
                for category, skill_list in skills.items():
                    if not isinstance(skill_list, list):
                        errors.append(f"Skills category '{category}' must be a list, got {type(skill_list).__name__}")
                    else:
                        for i, skill in enumerate(skill_list):
                            if not isinstance(skill, str):
                                errors.append(f"Skill {i+1} in category '{category}' must be a string, got {type(skill).__name__}: {skill}")
        
        # Validate additional_sections structure
        if 'additional_sections' in self.data:
            additional = self.data['additional_sections']
            if not isinstance(additional, dict):
                errors.append("'additional_sections' must be a dictionary")
            else:
                for section_name, section_data in additional.items():
                    if section_name in ['volunteer', 'publications']:
                        if not isinstance(section_data, list):
                            errors.append(f"'{section_name}' in additional_sections must be a list, got {type(section_data).__name__}")
                            if isinstance(section_data, str) and section_data.lower() == 'pass':
                                errors.append(f"'{section_name}' contains 'pass' - use empty list [] instead")
                        else:
                            for i, item in enumerate(section_data):
                                if not isinstance(item, dict):
                                    errors.append(f"'{section_name}' entry {i+1} must be a dictionary, got {type(item).__name__}: {item}")
        
        # Report errors
        if errors:
            logger.error("YAML structure validation failed:")
            for error in errors:
                logger.error(f"  - {error}")
            logger.error("\\nPlease fix these issues in your YAML file and try again.")
            logger.error("\\nFor help with YAML structure, see data/example_cv.yaml")
            return False
        
        return True
    
    def load_data(self) -> bool:
        """
        Load and validate YAML data.
        
        Returns:
            bool: True if data loaded successfully, False otherwise
        """
        try:
            if not os.path.exists(self.yaml_file):
                logger.error(f"YAML file not found: {self.yaml_file}")
                logger.info("Please copy data/example_cv.yaml to data/master_cv.yaml and update with your information")
                return False
            
            with open(self.yaml_file, 'r', encoding='utf-8') as file:
                self.data = yaml.safe_load(file)
            
            if not self.data:
                logger.error("YAML file is empty or invalid")
                return False
            
            # Extract configuration
            self.config = self.data.get('cv_config', {})
            
            # Validate YAML structure
            if not self.validate_yaml_structure():
                return False
            
            logger.info("YAML data loaded successfully")
            return True
            
        except yaml.YAMLError as e:
            logger.error(f"Error parsing YAML file: {e}")
            logger.error("Please check your YAML syntax. Common issues:")
            logger.error("  - Incorrect indentation (use spaces, not tabs)")
            logger.error("  - Missing colons after keys")
            logger.error("  - Unquoted strings with special characters")
            logger.error("  - Using 'pass' instead of empty lists []")
            return False
        except Exception as e:
            logger.error(f"Error loading data: {e}")
            return False
    
    def create_document(self):
        """Create a new Word document with ATS-friendly formatting."""
        self.doc = Document()
        
        # Set document margins (1 inch on all sides for ATS compatibility)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        logger.debug("Created new document with ATS-friendly margins")
    
    def add_personal_info(self):
        """Add personal information section to the document."""
        if not self.personal_info:
            logger.warning("No personal information found in environment variables")
            return
        
        personal = self.personal_info
        
        # Add name as main heading
        name_para = self.doc.add_heading(personal.get('name', ''), level=1)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact information - each on separate line for better ATS parsing
        if personal.get('email'):
            email_para = self.doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add hyperlink for email
            self._add_hyperlink_simple(email_para, personal['email'], f"mailto:{personal['email']}")
            
        if personal.get('phone'):
            phone_para = self.doc.add_paragraph()
            phone_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add hyperlink for phone
            self._add_hyperlink_simple(phone_para, personal['phone'], f"tel:{personal['phone']}")
            
        if personal.get('location'):
            location_para = self.doc.add_paragraph(personal['location'])
            location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if personal.get('linkedin'):
            linkedin_para = self.doc.add_paragraph()
            linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add hyperlink for LinkedIn
            self._add_hyperlink_simple(linkedin_para, f"LinkedIn: {personal['linkedin']}", personal['linkedin'])
            
        if personal.get('website'):
            website_para = self.doc.add_paragraph()
            website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add hyperlink for website
            self._add_hyperlink_simple(website_para, f"Website: {personal['website']}", personal['website'])
            
        if personal.get('github'):
            github_para = self.doc.add_paragraph()
            github_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add hyperlink for GitHub
            self._add_hyperlink_simple(github_para, f"GitHub: {personal['github']}", personal['github'])
        
        logger.info(f"Added personal information for: {personal.get('name', 'Unknown')}")
    
    def _add_hyperlink_simple(self, paragraph, text: str, url: str):
        """
        Add a simple hyperlink to a paragraph using a more reliable method.
        
        Args:
            paragraph: The paragraph object to add the hyperlink to
            text: The display text for the hyperlink
            url: The URL to link to
        """
        try:
            # For now, just add the text with proper formatting
            # TODO: Implement proper hyperlink functionality
            run = paragraph.add_run(text)
            run.font.color.rgb = None  # Use default color
            run.font.underline = True
            
            # Log the hyperlink for debugging
            logger.debug(f"Added hyperlink: {text} -> {url}")
            
        except Exception as e:
            logger.warning(f"Could not add hyperlink for {text}: {e}")
            # Fallback: just add the text without hyperlink
            paragraph.add_run(text)
    
    def _add_hyperlink(self, paragraph, text: str, url: str):
        """
        Add a hyperlink to a paragraph in the DOCX document.
        
        Args:
            paragraph: The paragraph object to add the hyperlink to
            text: The display text for the hyperlink
            url: The URL to link to
        """
        try:
            # Clear the paragraph first
            paragraph.clear()
            
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), url)
            
            # Create run element
            run = OxmlElement('w:r')
            run_props = OxmlElement('w:rPr')
            
            # Set hyperlink styling (blue color, underlined)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')
            run_props.append(color)
            
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            run_props.append(underline)
            
            run.append(run_props)
            
            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            
        except Exception as e:
            logger.warning(f"Could not add hyperlink for {text}: {e}")
            # Fallback: just add the text without hyperlink
            paragraph.add_run(text)
    
    def add_summary(self):
        """Add professional summary section."""
        if 'summary' not in self.data:
            logger.debug("No summary found in YAML data")
            return
        
        summary_text = self.data['summary'].strip()
        if not summary_text:
            return
        
        self.doc.add_heading('Professional Summary', level=2)
        summary_para = self.doc.add_paragraph(summary_text)
        # Add small spacing after summary
        self.doc.add_paragraph()
        
        logger.info("Added professional summary section")
    
    def add_experience(self):
        """Add work experience section."""
        if 'experience' not in self.data:
            logger.debug("No experience found in YAML data")
            return
        
        experiences = self.data['experience']
        if not experiences:
            return
        
        self.doc.add_heading('Professional Experience', level=2)
        
        try:
            for i, exp in enumerate(experiences):
                # Check if exp is a dictionary
                if not isinstance(exp, dict):
                    logger.error(f"Experience entry {i+1} must be a dictionary, got {type(exp).__name__}: {exp}")
                    logger.error("Expected format: - company: 'Company Name'\\n  role: 'Job Title'\\n  start_date: 'YYYY-MM'\\n  end_date: 'YYYY-MM' or 'Present'")
                    continue
                
                # Job title and company
                title_company = f"{exp.get('role', '')} - {exp.get('company', '')}"
                if exp.get('location'):
                    title_company += f" ({exp['location']})"
                
                self.doc.add_heading(title_company, level=3)
                
                # Date range
                start_date = exp.get('start_date', '')
                end_date = exp.get('end_date', 'Present')
                date_range = f"{start_date} - {end_date}"
                
                date_para = self.doc.add_paragraph(date_range)
                date_para.paragraph_format.italic = True
                
                # Description
                if exp.get('description'):
                    desc_para = self.doc.add_paragraph(exp['description'])
                
                # Achievements
                if exp.get('achievements'):
                    if not isinstance(exp['achievements'], list):
                        logger.error(f"Experience {i+1}: 'achievements' must be a list, got {type(exp['achievements']).__name__}")
                        logger.error("Expected format: achievements: ['Achievement 1', 'Achievement 2']")
                    else:
                        for achievement in exp['achievements']:
                            if not isinstance(achievement, str):
                                logger.error(f"Experience {i+1}: Each achievement must be a string, got {type(achievement).__name__}: {achievement}")
                                continue
                            achievement_para = self.doc.add_paragraph(achievement, style='List Bullet')
                
                # Technologies
                if exp.get('technologies'):
                    if not isinstance(exp['technologies'], list):
                        logger.error(f"Experience {i+1}: 'technologies' must be a list, got {type(exp['technologies']).__name__}")
                        logger.error("Expected format: technologies: ['Technology 1', 'Technology 2']")
                    else:
                        tech_text = f"Technologies: {', '.join(exp['technologies'])}"
                        tech_para = self.doc.add_paragraph(tech_text)
                        tech_para.paragraph_format.italic = True
                
                # Add spacing between jobs
                self.doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error processing experience section: {e}")
            logger.error("Make sure experience entries are properly formatted as dictionaries")
            return
        
        logger.info(f"Added {len(experiences)} work experience entries")
    
    def add_education(self):
        """Add education section."""
        if 'education' not in self.data:
            logger.debug("No education found in YAML data")
            return
        
        education_list = self.data['education']
        if not education_list:
            return
        
        self.doc.add_heading('Education', level=2)
        
        try:
            for i, edu in enumerate(education_list):
                # Check if edu is a dictionary
                if not isinstance(edu, dict):
                    logger.error(f"Education entry {i+1} must be a dictionary, got {type(edu).__name__}: {edu}")
                    logger.error("Expected format: - degree: 'Degree Name'\\n  institution: 'Institution Name'\\n  graduation_date: 'YYYY-MM'")
                    continue
                
                # Degree and institution
                degree_inst = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                if edu.get('location'):
                    degree_inst += f" ({edu['location']})"
                
                self.doc.add_heading(degree_inst, level=3)
                
                # Graduation date
                grad_date = edu.get('graduation_date', '')
                if grad_date:
                    date_para = self.doc.add_paragraph(f"Graduated: {grad_date}")
                    date_para.paragraph_format.italic = True
                
                # GPA (if provided)
                if edu.get('gpa'):
                    gpa_para = self.doc.add_paragraph(f"GPA: {edu['gpa']}")
                
                # Honors (if provided)
                if edu.get('honors'):
                    honors_para = self.doc.add_paragraph(f"Honors: {edu['honors']}")
                
                # Relevant coursework
                if edu.get('relevant_coursework'):
                    if not isinstance(edu['relevant_coursework'], list):
                        logger.error(f"Education {i+1}: 'relevant_coursework' must be a list, got {type(edu['relevant_coursework']).__name__}")
                        logger.error("Expected format: relevant_coursework: ['Course 1', 'Course 2']")
                    else:
                        coursework_text = f"Relevant Coursework: {', '.join(edu['relevant_coursework'])}"
                        coursework_para = self.doc.add_paragraph(coursework_text)
                
                self.doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error processing education section: {e}")
            logger.error("Make sure education entries are properly formatted as dictionaries")
            return
        
        logger.info(f"Added {len(education_list)} education entries")
    
    def add_skills(self):
        """Add skills section."""
        if 'skills' not in self.data:
            logger.debug("No skills found in YAML data")
            return
        
        skills = self.data['skills']
        if not skills:
            return
        
        self.doc.add_heading('Skills', level=2)
        
        for category, skill_list in skills.items():
            if skill_list:
                # Format category name (replace underscores with spaces, title case)
                category_name = category.replace('_', ' ').title()
                self.doc.add_heading(category_name, level=3)
                
                # Add skills as comma-separated list
                skills_text = ', '.join(skill_list)
                skills_para = self.doc.add_paragraph(skills_text)
        
        logger.info("Added skills section")
    
    def add_certifications(self):
        """Add certifications section."""
        if 'certifications' not in self.data:
            logger.debug("No certifications found in YAML data")
            return
        
        certs = self.data['certifications']
        if not certs or certs is None:
            return
        
        self.doc.add_heading('Certifications', level=2)
        
        try:
            for i, cert in enumerate(certs):
                # Check if cert is a dictionary
                if not isinstance(cert, dict):
                    logger.error(f"Certification entry {i+1} must be a dictionary, got {type(cert).__name__}: {cert}")
                    logger.error("Expected format: - name: 'Certification Name'\\n  issuer: 'Issuing Organization'\\n  date: 'YYYY-MM'")
                    continue
                
                cert_text = f"{cert.get('name', '')} - {cert.get('issuer', '')}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                
                cert_para = self.doc.add_paragraph(cert_text)
                
                if cert.get('credential_id'):
                    id_para = self.doc.add_paragraph(f"Credential ID: {cert['credential_id']}")
                    id_para.paragraph_format.italic = True
        except Exception as e:
            logger.error(f"Error processing certifications section: {e}")
            logger.error("Make sure certification entries are properly formatted as dictionaries")
            return
        
        logger.info(f"Added {len(certs)} certifications")
    
    def add_projects(self):
        """Add projects section."""
        if 'projects' not in self.data:
            logger.debug("No projects found in YAML data")
            return
        
        projects = self.data['projects']
        if not projects:
            return
        
        self.doc.add_heading('Projects', level=2)
        
        try:
            for i, project in enumerate(projects):
                # Check if project is a dictionary
                if not isinstance(project, dict):
                    logger.error(f"Project entry {i+1} must be a dictionary, got {type(project).__name__}: {project}")
                    logger.error("Expected format: - name: 'Project Name'\\n  description: 'Description'\\n  technologies: ['Tech 1', 'Tech 2']")
                    continue
                
                # Project name
                self.doc.add_heading(project.get('name', ''), level=3)
                
                # Description
                if project.get('description'):
                    desc_para = self.doc.add_paragraph(project['description'])
                
                # Technologies
                if project.get('technologies'):
                    if not isinstance(project['technologies'], list):
                        logger.error(f"Project {i+1}: 'technologies' must be a list, got {type(project['technologies']).__name__}")
                        logger.error("Expected format: technologies: ['Technology 1', 'Technology 2']")
                    else:
                        tech_text = f"Technologies: {', '.join(project['technologies'])}"
                        tech_para = self.doc.add_paragraph(tech_text)
                        tech_para.paragraph_format.italic = True
                
                # URL (if provided)
                if project.get('url'):
                    url_para = self.doc.add_paragraph()
                    # Add hyperlink for project URL
                    self._add_hyperlink_simple(url_para, f"URL: {project['url']}", project['url'])
                
                # Date (if provided)
                if project.get('date'):
                    date_para = self.doc.add_paragraph(f"Date: {project['date']}")
                
                self.doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error processing projects section: {e}")
            logger.error("Make sure project entries are properly formatted as dictionaries")
            return
        
        logger.info(f"Added {len(projects)} projects")
    
    def add_languages(self):
        """Add languages section."""
        if 'languages' not in self.data:
            logger.debug("No languages found in YAML data")
            return
        
        languages = self.data['languages']
        if not languages:
            return
        
        self.doc.add_heading('Languages', level=2)
        
        try:
            for lang in languages:
                # Check if lang is a dictionary
                if not isinstance(lang, dict):
                    logger.error(f"Language entry must be a dictionary, got {type(lang).__name__}: {lang}")
                    logger.error("Expected format: - language: 'Language Name'\\n  proficiency: 'Proficiency Level'")
                    continue
                
                lang_text = f"{lang.get('language', '')} - {lang.get('proficiency', '')}"
                lang_para = self.doc.add_paragraph(lang_text)
        except Exception as e:
            logger.error(f"Error processing languages section: {e}")
            logger.error("Make sure language entries are properly formatted as dictionaries")
            return
        
        logger.info(f"Added {len(languages)} languages")
    
    def add_additional_sections(self):
        """Add additional sections like volunteer work, publications, etc."""
        if 'additional_sections' not in self.data:
            logger.debug("No additional sections found in YAML data")
            return
        
        additional = self.data['additional_sections']
        
        # Volunteer work
        if 'volunteer' in additional and additional['volunteer']:
            self.doc.add_heading('Volunteer Experience', level=2)
            try:
                for vol in additional['volunteer']:
                    # Check if vol is a dictionary
                    if not isinstance(vol, dict):
                        logger.error(f"Volunteer entry must be a dictionary, got {type(vol).__name__}: {vol}")
                        logger.error("Expected format: - role: 'Role Name'\\n  organization: 'Organization'\\n  duration: 'Duration'")
                        continue
                    
                    vol_text = f"{vol.get('role', '')} - {vol.get('organization', '')}"
                    if vol.get('duration'):
                        vol_text += f" ({vol['duration']})"
                    
                    self.doc.add_heading(vol_text, level=3)
                    if vol.get('description'):
                        desc_para = self.doc.add_paragraph(vol['description'])
                    self.doc.add_paragraph()
            except Exception as e:
                logger.error(f"Error processing volunteer section: {e}")
                logger.error("Make sure volunteer entries are properly formatted as dictionaries")
        
        # Publications
        if 'publications' in additional and additional['publications']:
            self.doc.add_heading('Publications', level=2)
            try:
                for pub in additional['publications']:
                    # Check if pub is a dictionary
                    if not isinstance(pub, dict):
                        logger.error(f"Publication entry must be a dictionary, got {type(pub).__name__}: {pub}")
                        logger.error("Expected format: - title: 'Title'\\n  publication: 'Publication Name'\\n  date: 'Date'")
                        continue
                    
                    pub_text = f"{pub.get('title', '')} - {pub.get('publication', '')}"
                    if pub.get('date'):
                        pub_text += f" ({pub['date']})"
                    
                    self.doc.add_heading(pub_text, level=3)
                    if pub.get('url'):
                        url_para = self.doc.add_paragraph()
                        # Add hyperlink for publication URL
                        self._add_hyperlink_simple(url_para, f"URL: {pub['url']}", pub['url'])
                    self.doc.add_paragraph()
            except Exception as e:
                logger.error(f"Error processing publications section: {e}")
                logger.error("Make sure publication entries are properly formatted as dictionaries")
        
        logger.info("Added additional sections")
    
    def add_secret_message(self):
        """
        Add a hidden secret message with white text for AI detection systems.
        The text is invisible to human readers but can be detected by AI scanning systems.
        """
        # Get secret message from config or data
        secret_message = self.config.get('secret_message', self.data.get('secret_message', ''))
        
        if not secret_message:
            logger.debug("No secret message found in YAML data")
            return
        
        logger.info("Adding secret message (white text)")
        
        # Add a paragraph with white text
        secret_para = self.doc.add_paragraph()
        secret_run = secret_para.add_run(secret_message)
        
        # Set font color to white (RGB 255, 255, 255) - invisible on white background
        secret_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Use same font as the rest of the document for consistency
        font_family = self.config.get('font_family', 'Arial')
        font_size = self.config.get('font_size', 11)
        secret_run.font.name = font_family
        secret_run.font.size = Pt(font_size)
        
        # Make it very small to avoid layout issues (optional, but helps)
        secret_run.font.size = Pt(1)
        
        # Add minimal spacing
        secret_para.paragraph_format.space_after = Pt(0)
        secret_para.paragraph_format.space_before = Pt(0)
        
        logger.debug(f"Secret message added: {len(secret_message)} characters")
    
    def apply_formatting(self):
        """Apply ATS-friendly formatting to the document."""
        if not self.doc:
            return
        
        # Get font settings from config
        font_family = self.config.get('font_family', 'Arial')
        font_size = self.config.get('font_size', 11)
        
        # Apply formatting to all paragraphs
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_family
                run.font.size = Pt(font_size)
            
            # Add proper spacing after headings
            if paragraph.style.name.startswith('Heading'):
                paragraph.paragraph_format.space_after = Pt(6)
            else:
                paragraph.paragraph_format.space_after = Pt(3)
        
        logger.debug(f"Applied formatting: {font_family}, {font_size}pt")
    
    def convert_docx_to_pdf(self, docx_path: str) -> str:
        """
        Convert DOCX file to PDF using available PDF generation libraries.
        
        Args:
            docx_path: Path to the DOCX file to convert
            
        Returns:
            str: Path to the generated PDF file, or empty string if conversion failed
        """
        if not os.path.exists(docx_path):
            logger.error(f"DOCX file not found: {docx_path}")
            return ""
        
        # Generate PDF filename
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Try WeasyPrint first (better formatting)
            if WEASYPRINT_AVAILABLE:
                logger.info("Converting DOCX to PDF using WeasyPrint...")
                return self._convert_with_weasyprint(docx_path, pdf_path)
            
            # Fallback to pdfkit if WeasyPrint not available
            elif PDFKIT_AVAILABLE:
                logger.info("Converting DOCX to PDF using pdfkit...")
                return self._convert_with_pdfkit(docx_path, pdf_path)
            
            else:
                logger.error("No PDF generation libraries available. Please install weasyprint or pdfkit.")
                return ""
                
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            return ""
    
    def _convert_with_weasyprint(self, docx_path: str, pdf_path: str) -> str:
        """
        Convert DOCX to PDF using WeasyPrint.
        This method converts the DOCX to HTML first, then to PDF.
        """
        try:
            # For now, we'll use a simple approach: convert DOCX to HTML then to PDF
            # This is a basic implementation - in production you might want to use
            # python-docx2txt or mammoth for better HTML conversion
            logger.warning("WeasyPrint conversion requires HTML input. Using basic text extraction.")
            
            # Extract text content from DOCX (simplified approach)
            doc = Document(docx_path)
            html_content = self._docx_to_html(doc)
            
            # Convert HTML to PDF using WeasyPrint
            html_doc = weasyprint.HTML(string=html_content)
            html_doc.write_pdf(pdf_path)
            
            logger.info(f"PDF generated successfully: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"WeasyPrint conversion failed: {e}")
            return ""
    
    def _convert_with_pdfkit(self, docx_path: str, pdf_path: str) -> str:
        """
        Convert DOCX to PDF using pdfkit.
        This method requires wkhtmltopdf to be installed on the system.
        """
        try:
            # pdfkit requires HTML input, so we need to convert DOCX to HTML first
            logger.warning("pdfkit conversion requires HTML input. Using basic text extraction.")
            
            # Extract text content from DOCX (simplified approach)
            doc = Document(docx_path)
            html_content = self._docx_to_html(doc)
            
            # Configure pdfkit options
            options = {
                'page-size': 'A4',
                'margin-top': '1in',
                'margin-right': '1in',
                'margin-bottom': '1in',
                'margin-left': '1in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            # Convert HTML to PDF
            pdfkit.from_string(html_content, pdf_path, options=options)
            
            logger.info(f"PDF generated successfully: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"pdfkit conversion failed: {e}")
            return ""
    
    def _docx_to_html(self, doc: Document) -> str:
        """
        Convert a python-docx Document to HTML with enhanced styling.
        This implementation provides better formatting for PDF generation.
        """
        html_parts = []
        html_parts.append("<!DOCTYPE html>")
        html_parts.append("<html><head>")
        html_parts.append("<meta charset='UTF-8'>")
        html_parts.append("<style>")
        html_parts.append("""
        body { 
            font-family: 'Arial', 'Helvetica', sans-serif; 
            font-size: 11pt; 
            line-height: 1.5; 
            margin: 0.8in; 
            color: #333;
            background-color: #ffffff;
        }
        h1 { 
            font-size: 24pt; 
            font-weight: bold; 
            text-align: center; 
            margin-bottom: 20pt; 
            margin-top: 0;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10pt;
        }
        h2 { 
            font-size: 16pt; 
            font-weight: bold; 
            margin-top: 24pt; 
            margin-bottom: 12pt; 
            color: #2c3e50;
            border-left: 4px solid #3498db;
            padding-left: 10pt;
        }
        h3 { 
            font-size: 13pt; 
            font-weight: bold; 
            margin-top: 16pt; 
            margin-bottom: 8pt; 
            color: #34495e;
        }
        p { 
            margin-bottom: 8pt; 
            text-align: justify;
        }
        ul { 
            margin-bottom: 12pt; 
            padding-left: 20pt;
        }
        li { 
            margin-bottom: 4pt; 
            line-height: 1.4;
        }
        .contact-info {
            text-align: center;
            margin-bottom: 20pt;
            font-size: 10pt;
            color: #7f8c8d;
        }
        .contact-info a {
            color: #3498db;
            text-decoration: none;
        }
        .contact-info a:hover {
            text-decoration: underline;
        }
        .date-range {
            font-style: italic;
            color: #7f8c8d;
            font-size: 10pt;
        }
        .company-role {
            font-weight: bold;
            color: #2c3e50;
        }
        .achievement {
            margin-left: 15pt;
            position: relative;
        }
        .achievement:before {
            content: "•";
            color: #3498db;
            font-weight: bold;
            position: absolute;
            left: -15pt;
        }
        ul {
            list-style-type: none;
            padding-left: 0;
        }
        .secret-message {
            color: #ffffff !important;
            font-size: 1pt;
            margin: 0;
            padding: 0;
            line-height: 0;
            opacity: 1;
        }
        """)
        html_parts.append("</style>")
        html_parts.append("</head><body>")
        
        # Track if we're in a list
        in_list = False
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check if this paragraph contains white text (secret message)
                is_white_text = self._is_white_text(paragraph)
                
                # Determine if this is a heading based on style
                if paragraph.style.name.startswith('Heading 1'):
                    if in_list:
                        html_parts.append("</ul>")
                        in_list = False
                    html_parts.append(f"<h1>{self._make_links_clickable(paragraph.text)}</h1>")
                elif paragraph.style.name.startswith('Heading 2'):
                    if in_list:
                        html_parts.append("</ul>")
                        in_list = False
                    html_parts.append(f"<h2>{self._make_links_clickable(paragraph.text)}</h2>")
                elif paragraph.style.name.startswith('Heading 3'):
                    if in_list:
                        html_parts.append("</ul>")
                        in_list = False
                    html_parts.append(f"<h3>{self._make_links_clickable(paragraph.text)}</h3>")
                else:
                    # Check if it's a bullet point
                    if paragraph.style.name == 'List Bullet':
                        if not in_list:
                            html_parts.append("<ul>")
                            in_list = True
                        html_parts.append(f"<li class='achievement'>{self._make_links_clickable(paragraph.text)}</li>")
                    else:
                        if in_list:
                            html_parts.append("</ul>")
                            in_list = False
                        # Check if this is white text (secret message)
                        if is_white_text:
                            html_parts.append(f"<p class='secret-message'>{paragraph.text}</p>")
                        # Check if this looks like a date range
                        elif self._is_date_range(paragraph.text):
                            html_parts.append(f"<p class='date-range'>{paragraph.text}</p>")
                        else:
                            html_parts.append(f"<p>{self._make_links_clickable(paragraph.text)}</p>")
        
        # Close any remaining list
        if in_list:
            html_parts.append("</ul>")
        
        html_parts.append("</body></html>")
        return "\n".join(html_parts)
    
    def _make_links_clickable(self, text: str) -> str:
        """
        Convert email addresses, phone numbers, and URLs to clickable links.
        """
        import re
        
        # Email addresses
        text = re.sub(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', 
                      r'<a href="mailto:\1">\1</a>', text)
        
        # Phone numbers (various formats) - be more specific to avoid false matches
        text = re.sub(r'(\+?[\d\s\-\(\)]{10,})', 
                      r'<a href="tel:\1">\1</a>', text)
        
        # URLs (http/https) - be more specific
        text = re.sub(r'(https?://[^\s<>"]+)', 
                      r'<a href="\1" target="_blank">\1</a>', text)
        
        # LinkedIn URLs - handle both with and without https
        text = re.sub(r'LinkedIn:\s*(https?://[^\s<>"]+)', 
                      r'LinkedIn: <a href="\1" target="_blank">\1</a>', text)
        text = re.sub(r'LinkedIn:\s*([^\s<>"]+)', 
                      r'LinkedIn: <a href="https://\1" target="_blank">\1</a>', text)
        
        # GitHub URLs - handle both with and without https
        text = re.sub(r'GitHub:\s*(https?://[^\s<>"]+)', 
                      r'GitHub: <a href="\1" target="_blank">\1</a>', text)
        text = re.sub(r'GitHub:\s*([^\s<>"]+)', 
                      r'GitHub: <a href="https://\1" target="_blank">\1</a>', text)
        
        # Website URLs - handle both with and without https
        text = re.sub(r'Website:\s*(https?://[^\s<>"]+)', 
                      r'Website: <a href="\1" target="_blank">\1</a>', text)
        text = re.sub(r'Website:\s*([^\s<>"]+)', 
                      r'Website: <a href="https://\1" target="_blank">\1</a>', text)
        
        return text
    
    def _is_date_range(self, text: str) -> bool:
        """
        Check if text looks like a date range (e.g., "2024-10 - Present").
        """
        import re
        # Pattern for date ranges like "2024-10 - Present" or "2023-01 - 2024-12"
        date_pattern = r'\d{4}-\d{2}\s*-\s*(Present|\d{4}-\d{2})'
        return bool(re.match(date_pattern, text.strip()))
    
    def _is_white_text(self, paragraph) -> bool:
        """
        Check if a paragraph contains white text (used for secret messages).
        
        Args:
            paragraph: The paragraph object to check
            
        Returns:
            bool: True if the paragraph contains white text, False otherwise
        """
        try:
            # Check all runs in the paragraph
            for run in paragraph.runs:
                # Check if the run has a color set
                if run.font.color.rgb:
                    # Check if color is white (RGB 255, 255, 255)
                    if run.font.color.rgb == RGBColor(255, 255, 255):
                        return True
            return False
        except Exception as e:
            logger.debug(f"Error checking for white text: {e}")
            return False
    
    def generate_filename(self) -> str:
        """Generate filename for the CV document."""
        # Use filename prefix from environment variable if available
        filename_prefix = os.getenv('CV_FILENAME_PREFIX', '')
        if filename_prefix:
            name = filename_prefix.replace(' ', '_')
        else:
            # Fall back to name from personal info
            name = self.personal_info.get('name', 'CV').replace(' ', '_')
        
        # Get filename prefix from config
        prefix = self.config.get('filename_prefix', name)
        
        # Add timestamp if configured
        if self.config.get('include_timestamp', True):
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{prefix}_{timestamp}.docx"
        else:
            filename = f"{prefix}.docx"
        
        return filename
    
    def save_document(self) -> str:
        """Save the document to the output folder and generate PDF."""
        if not self.doc:
            logger.error("No document to save")
            return ""
        
        # Create output directory if it doesn't exist
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        
        # Generate filename
        filename = self.generate_filename()
        filepath = output_dir / filename
        
        try:
            # Save DOCX file
            self.doc.save(str(filepath))
            logger.info(f"CV saved successfully: {filepath}")
            
            # Generate PDF if PDF generation is available
            if WEASYPRINT_AVAILABLE or PDFKIT_AVAILABLE:
                logger.info("Generating PDF version...")
                pdf_path = self.convert_docx_to_pdf(str(filepath))
                if pdf_path:
                    logger.info(f"PDF generated successfully: {pdf_path}")
                else:
                    logger.warning("PDF generation failed, but DOCX was saved successfully")
            else:
                logger.info("PDF generation libraries not available. Only DOCX file saved.")
            
            return str(filepath)
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            return ""
    
    def generate_cv(self) -> bool:
        """
        Generate the complete CV document.
        
        Returns:
            bool: True if generation was successful, False otherwise
        """
        try:
            logger.info("Starting CV generation process")
            
            # Load data
            if not self.load_data():
                return False
            
            # Create document
            self.create_document()
            
            # Add sections in order
            section_order = self.config.get('section_order', [
                'personal_info', 'summary', 'experience', 'education', 
                'skills', 'certifications', 'projects', 'languages', 'additional_sections'
            ])
            
            # Map section names to methods
            section_methods = {
                'personal_info': self.add_personal_info,
                'summary': self.add_summary,
                'experience': self.add_experience,
                'education': self.add_education,
                'skills': self.add_skills,
                'certifications': self.add_certifications,
                'projects': self.add_projects,
                'languages': self.add_languages,
                'additional_sections': self.add_additional_sections
            }
            
            # Add sections in configured order
            for section in section_order:
                if section in section_methods and section not in self.config.get('hidden_sections', []):
                    section_methods[section]()
            
            # Add secret message (white text for AI detection) at the end
            self.add_secret_message()
            
            # Apply formatting
            self.apply_formatting()
            
            # Save document
            filepath = self.save_document()
            if not filepath:
                return False
            
            logger.info("CV generation completed successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error during CV generation: {e}")
            return False


def main():
    """Main function to run the CV generator."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate CV from YAML data')
    parser.add_argument('--yaml-file', default='data/master_cv.yaml',
                       help='Path to YAML file containing CV data')
    parser.add_argument('--output-dir', default='output',
                       help='Output directory for generated CV')
    
    args = parser.parse_args()
    
    # Initialize generator
    generator = CVGenerator(args.yaml_file)
    
    # Generate CV
    success = generator.generate_cv()
    
    if success:
        logger.info("CV generation completed successfully!")
        sys.exit(0)
    else:
        logger.error("CV generation failed!")
        sys.exit(1)


if __name__ == "__main__":
    main()
